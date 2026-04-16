"""
LangGraph API Routes - RESTful endpoints for the simplified pipeline.
Provides FastAPI integration for the 4-agent LangGraph orchestration.
"""

from fastapi import APIRouter, HTTPException
from fastapi.responses import FileResponse
from pathlib import Path
from pydantic import BaseModel
from typing import List, Optional, Dict, Any
import json
import os
import re
import math
import tempfile
from dotenv import load_dotenv

# Load environment variables immediately
load_dotenv()

from backend.services.graph_integration import GraphIntegrationService

# Ensure outputs directory always exists at startup
Path("outputs").mkdir(parents=True, exist_ok=True)


def _safe_filename(value: str, fallback: str = "paper") -> str:
    """Return a filesystem-safe filename stem."""
    base = (value or fallback).strip()
    sanitized = re.sub(r"[^A-Za-z0-9._-]+", "_", base)
    sanitized = sanitized.strip("._")
    return sanitized or fallback



def _build_fallback_docx(result: Dict[str, Any], task_id: str) -> str:
    """Generate a DOCX file from paper sections using python-docx."""
    try:
        from docx import Document  # type: ignore

        doc = Document()
        title = result.get("topic") or "Research Paper"
        doc.add_heading(title, level=0)

        for heading, key in [
            ("Abstract",              "abstract"),
            ("Introduction",          "introduction"),
            ("Related Work",          "related_work"),
            ("Methodology",           "methodology"),
            ("Implementation",        "implementation"),
            ("Results and Discussion", "results_discussion"),
            ("Conclusion",            "conclusion"),
        ]:
            body = result.get(key, "") or result.get("results", "") if key == "results_discussion" else result.get(key, "")
            if body:
                doc.add_heading(heading, level=1)
                doc.add_paragraph(str(body))

        refs = result.get("references_raw", []) or result.get("references", [])
        if refs:
            doc.add_heading("References", level=1)
            for i, ref in enumerate(refs, 1):
                label = ref.get("title", str(ref)) if isinstance(ref, dict) else str(ref)
                doc.add_paragraph(f"[{i}] {label}")

        Path("outputs").mkdir(parents=True, exist_ok=True)
        out_path = f"outputs/{task_id}.docx"
        doc.save(out_path)
        return out_path
    except Exception as exc:
        raise RuntimeError(f"DOCX generation failed: {exc}") from exc


def _get_task_result(task_id: str) -> Dict[str, Any]:
    """
    Shared helper: validate task exists and is completed, then return its result dict.
    Raises HTTPException 404 with a specific message for each failure mode.
    """
    task = graph_service.get_task(task_id)
    if not task:
        raise HTTPException(status_code=404, detail=f"Task {task_id} not found")
    if task.get("status") != "completed":
        raise HTTPException(status_code=404, detail=f"Task {task_id} has not completed yet (status: {task.get('status', 'unknown')})")
    result = graph_service.get_result(task_id) or {}
    return result


def _sections_from_result(result: Dict[str, Any]) -> Dict[str, str]:
    """Build sections map expected by IEEE formatter from task result fields."""
    return {
        "Introduction": result.get("introduction", ""),
        "Related Work": result.get("related_work", ""),
        "Methodology": result.get("methodology", ""),
        "Implementation": result.get("implementation", ""),
        "Results": result.get("results", ""),
        "Conclusion": result.get("conclusion", ""),
    }


def _authors_from_result(result: Dict[str, Any]) -> List[Dict[str, str]]:
    """Normalize author payload for IEEE formatter."""
    authors = result.get("authors", [])
    if isinstance(authors, list):
        normalized: List[Dict[str, str]] = []
        for item in authors:
            if isinstance(item, dict):
                normalized.append(item)
            elif isinstance(item, str) and item.strip():
                normalized.append({"name": item.strip(), "affiliation": "", "location": "", "email": ""})
        return normalized
    if isinstance(authors, str) and authors.strip():
        return [{"name": authors.strip(), "affiliation": "", "location": "", "email": ""}]
    return []


def _estimate_syllables(word: str) -> int:
    """Rough syllable estimator for readability scoring."""
    w = re.sub(r"[^a-z]", "", word.lower())
    if not w:
        return 1
    groups = re.findall(r"[aeiouy]+", w)
    count = max(1, len(groups))
    if w.endswith("e") and count > 1:
        count -= 1
    return max(1, count)


def _build_validation_response(result: Dict[str, Any]) -> Dict[str, Any]:
    """Build dashboard validation payload from generated paper sections."""
    sections = [
        result.get("abstract", ""),
        result.get("introduction", ""),
        result.get("related_work", ""),
        result.get("methodology", ""),
        result.get("implementation", ""),
        result.get("results", ""),
        result.get("conclusion", ""),
    ]
    text = "\n\n".join([s for s in sections if isinstance(s, str) and s.strip()])

    words = re.findall(r"\b\w+\b", text)
    word_count = len(words)
    sentence_parts = re.split(r"[.!?]+", text)
    sentence_count = len([s for s in sentence_parts if s.strip()])
    avg_sentence_length = round(word_count / max(sentence_count, 1), 2)
    unique_words = len(set([w.lower() for w in words]))
    vocabulary_richness = round(unique_words / max(word_count, 1), 3)

    syllables = sum(_estimate_syllables(w) for w in words)
    flesch_kincaid = round(
        206.835 - 1.015 * (word_count / max(sentence_count, 1)) - 84.6 * (syllables / max(word_count, 1)),
        2,
    )

    has_intro = bool(result.get("introduction"))
    has_conclusion = bool(result.get("conclusion"))
    section_count = sum(1 for s in sections if isinstance(s, str) and s.strip())
    has_sections = section_count >= 4
    has_citations = bool(result.get("references"))

    repetition_ratio = round(max(0.0, 1 - vocabulary_richness), 3)
    topic_relevance_score = 0.9 if word_count > 120 else 0.7

    base_score = 50
    base_score += 10 if has_intro else 0
    base_score += 10 if has_conclusion else 0
    base_score += min(20, section_count * 3)
    base_score += 10 if has_citations else 0
    base_score += int(max(0, min(10, vocabulary_richness * 20)))
    base_score = max(0, min(100, base_score))

    quality_level = "Excellent" if base_score >= 85 else "Good" if base_score >= 70 else "Fair" if base_score >= 50 else "Poor"

    warnings = []
    if word_count < 180:
        warnings.append("Paper appears short; consider expanding section detail.")
    if not has_citations:
        warnings.append("No references found in generated output.")

    errors = []
    if not has_intro:
        errors.append("Missing introduction section content.")
    if not has_conclusion:
        errors.append("Missing conclusion section content.")

    return {
        "overall_quality_score": float(base_score),
        "quality_level": quality_level,
        "is_valid": len(errors) == 0,
        "content_metrics": {
            "word_count": word_count,
            "sentence_count": sentence_count,
            "avg_sentence_length": avg_sentence_length,
            "unique_words": unique_words,
            "vocabulary_richness": vocabulary_richness,
            "flesch_kincaid_grade": flesch_kincaid,
        },
        "structure_metrics": {
            "has_introduction": has_intro,
            "has_conclusion": has_conclusion,
            "has_sections": has_sections,
            "section_count": section_count,
            "has_citations": has_citations,
        },
        "quality_metrics": {
            "grammar_errors": 0,
            "spelling_errors": 0,
            "repetition_ratio": repetition_ratio,
            "topic_relevance_score": topic_relevance_score,
        },
        "validation_warnings": warnings,
        "validation_errors": errors,
    }

# Initialize router
router = APIRouter(prefix="/api/langgraph", tags=["langgraph"])

# Model startup check (Ollama-first)
model_provider = os.getenv("MODEL_PROVIDER", "ollama").strip().lower() or "ollama"
default_model = os.getenv("OLLAMA_MODEL", "llama3")

if model_provider != "ollama":
    print(f"[WARNING] MODEL_PROVIDER={model_provider} requested. Forcing ollama in ModelManager.")

print(f"[MODEL] Provider={model_provider} default_model={default_model}")

try:
    graph_service = GraphIntegrationService(model=default_model)
    print(f"[INFO] GraphIntegrationService initialized with {len(graph_service.tasks)} tasks loaded from disk")
except Exception as e:
    print(f"[ERROR] Failed to initialize GraphIntegrationService: {e}")
    graph_service = None


# Pydantic Models
class AuthorInfo(BaseModel):
    """Author information."""

    name:         str = "Author Name"
    department:   str = "Department"
    organization: str = "University"
    city_country: str = "City, Country"
    email:        str = "author@uni.edu"
    # kept for backward compat
    affiliation:  Optional[str] = None
    location:     Optional[str] = None


class EquationInfo(BaseModel):
    """Equation information."""

    id: Optional[str] = None
    label: str
    notation: str
    latex: Optional[str] = None
    description: Optional[str] = None


class DiagramInfo(BaseModel):
    """Diagram information."""

    id: Optional[str] = None
    label: str
    caption: str
    file: Optional[str] = None
    file_path: Optional[str] = None
    width: Optional[float] = None
    height: Optional[float] = None


class TableInfo(BaseModel):
    """Table information."""

    id: Optional[str] = None
    label: str
    caption: str
    headers: List[str] = []
    rows: List[List[str]] = []


class ResearchStructure(BaseModel):
    """Structured research information."""

    title: Optional[str] = None
    keywords: List[str] = []
    research_topic: Optional[str] = None
    problem_statement: Optional[str] = None
    proposed_solution: Optional[str] = None
    objective: Optional[str] = None
    methodology: Optional[str] = None
    dataset: Optional[str] = None
    experiments: Optional[str] = None
    equations: List[EquationInfo] = []
    diagrams: List[DiagramInfo] = []
    tables: List[TableInfo] = []
    notes: Optional[str] = None


class PaperGenerationRequest(BaseModel):
    """Request to generate a paper."""

    topic: str
    notes: Optional[str] = None
    diagrams: List[DiagramInfo] = []
    equations: List[EquationInfo] = []
    tables: List[TableInfo] = []
    authors: List[AuthorInfo] = []
    max_references: int = 5
    model_name: str = "llama3"
    research_structure: Optional[ResearchStructure] = None
    # Extended input schema
    paper_type: Optional[str] = None              # "conference" | "journal"
    target_venue: Optional[str] = None            # e.g. "IEEE CVPR 2025"
    page_limit: Optional[int] = None              # max pages
    problem_statement: Optional[str] = None       # clear problem description
    proposed_solution: Optional[str] = None       # high-level solution approach
    key_contributions: List[str] = []             # explicit contribution list
    datasets: List[str] = []                      # dataset names
    baselines: List[str] = []                     # comparison method names
    metrics: List[str] = []                       # evaluation metrics
    results_summary: Optional[str] = None         # expected/actual results context
    num_arxiv_papers: int = 10                    # number of arXiv papers to retrieve


class PaperGenerationResponse(BaseModel):
    """Response from paper generation."""

    task_id: str
    status: str
    message: str


class PaperStatusResponse(BaseModel):
    """Paper generation status."""

    task_id: str
    status: str
    step: str
    progress: int
    message: str
    stages: Dict[str, str]
    created_at: str
    started_at: Optional[str] = None
    completed_at: Optional[str] = None
    error: Optional[str] = None
    result: Optional[Dict[str, Any]] = None


class TaskListResponse(BaseModel):
    """List of tasks."""

    tasks: List[Dict[str, Any]]
    total: int


# Endpoints


@router.post("/generate", response_model=PaperGenerationResponse)
async def generate_paper(request: PaperGenerationRequest) -> PaperGenerationResponse:
    """
    Start paper generation using LangGraph pipeline.

    Args:
        request: Paper generation request with topic, diagrams, equations, etc.

    Returns:
        Task ID for status polling
    """
    try:
        if graph_service is None:
            raise HTTPException(
                status_code=503,
                detail="API service not initialized. Model backend is unavailable.",
            )
        
        # Convert Pydantic models to dicts
        diagrams = [d.dict() for d in request.diagrams]
        equations = [e.dict() for e in request.equations]
        tables = [t.dict() for t in request.tables]
        authors = [a.dict() for a in request.authors]
        research_structure = request.research_structure.dict() if request.research_structure else None

        # Start generation
        task_id = await graph_service.generate_paper(
            topic=request.topic,
            notes=request.notes,
            diagrams=diagrams,
            equations=equations,
            tables=tables,
            authors=authors,
            max_references=request.max_references,
            model_name=request.model_name,
            research_structure=research_structure,
            paper_type=request.paper_type,
            target_venue=request.target_venue,
            page_limit=request.page_limit,
            problem_statement=request.problem_statement,
            proposed_solution=request.proposed_solution,
            key_contributions=request.key_contributions,
            datasets=request.datasets,
            baselines=request.baselines,
            metrics=request.metrics,
            results_summary=request.results_summary,
            num_arxiv_papers=request.num_arxiv_papers,
        )
        
        print(f"[INFO] Task {task_id} created for topic: {request.topic}")
        print(f"[INFO] Tasks in service: {list(graph_service.tasks.keys())}")

        return PaperGenerationResponse(
            task_id=task_id,
            status="queued",
            message=f"Paper generation started. Task ID: {task_id}",
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Generation failed: {str(e)}")


@router.get("/status/{task_id}", response_model=PaperStatusResponse)
async def get_status(task_id: str) -> PaperStatusResponse:
    """
    Get paper generation status.

    Args:
        task_id: Task ID from generation endpoint

    Returns:
        Current status and progress
    """
    if graph_service is None:
        raise HTTPException(
            status_code=503,
            detail="API service not initialized. Model backend is unavailable.",
        )
    
    status = graph_service.get_status(task_id)

    # Only return 404 when the task does not exist.
    if status.get("error") == "Task not found":
        raise HTTPException(status_code=404, detail="Task not found")

    return PaperStatusResponse(**status)


@router.get("/validation/{task_id}")
async def get_validation(task_id: str):
    """Return validation metrics for a completed langgraph task."""
    if graph_service is None:
        raise HTTPException(
            status_code=503,
            detail="API service not initialized. Model backend is unavailable.",
        )

    result = _get_task_result(task_id)

    validation = _build_validation_response(result)
    if isinstance(result, dict) and result.get("ieee_report"):
        validation["ieee_report"] = result.get("ieee_report")
    if isinstance(result, dict) and result.get("enhancement_scores"):
        validation["enhancement_scores"] = result.get("enhancement_scores")
    return validation


@router.get("/download/{task_id}")
async def download_paper(task_id: str):
    """
    Download generated PDF.

    Resolution order:
    1. Serve cached PDF from task result if the file exists.
    2. Compile on demand: local pdflatex → latexonline.cc online API.
    3. Cache the compiled PDF path so subsequent calls skip recompilation.
    4. If both strategies fail, return HTTP 503 with JSON guidance.
    """
    if graph_service is None:
        raise HTTPException(
            status_code=503,
            detail="API service not initialized. Model backend is unavailable.",
        )

    result = _get_task_result(task_id)

    from backend.services.pdf_generator import PDFGenerator
    from fastapi.responses import JSONResponse

    # Helper to build the 503 JSON response
    def _compilation_failed_response() -> JSONResponse:
        return JSONResponse(
            status_code=503,
            content={
                "error": "PDF compilation failed",
                "suggestion": "Install pdflatex or download the LaTeX source and compile locally.",
                "latex_endpoint": f"/api/langgraph/export/latex/{task_id}",
                "install_pdflatex": {
                    "macOS": "brew install basictex",
                    "linux": "apt install texlive-latex-extra texlive-science",
                },
            },
        )

    # 1. Serve already-compiled PDF if FormattingAgent (or a prior download) produced one
    stored_pdf = result.get("pdf_path")
    if stored_pdf and Path(stored_pdf).exists():
        file_stem = _safe_filename(result.get("topic", "paper"))
        return FileResponse(
            path=stored_pdf,
            media_type="application/pdf",
            filename=f"{file_stem}_IEEE.pdf",
        )

    # 2. Compile on demand
    latex_source = result.get("latex_source", "")
    if not latex_source:
        # Try the saved .tex file in outputs/ as a last resort
        tex_candidate = Path("outputs") / f"{task_id}.tex"
        if tex_candidate.exists():
            latex_source = tex_candidate.read_text(encoding="utf-8")

    if not latex_source:
        return _compilation_failed_response()

    # generate_async: local pdflatex → latexonline.cc
    pdf_path = await PDFGenerator.generate_async(latex_source, task_id)

    if pdf_path and Path(pdf_path).exists():
        # 3. Cache so next download is instant
        graph_service.cache_pdf_path(task_id, pdf_path)
        file_stem = _safe_filename(result.get("topic", "paper"))
        return FileResponse(
            path=pdf_path,
            media_type="application/pdf",
            filename=f"{file_stem}_IEEE.pdf",
        )

    # 4. Both strategies failed
    return _compilation_failed_response()


@router.get("/latex/{task_id}")
async def get_latex(task_id: str):
    """
    Get LaTeX source code.

    Args:
        task_id: Task ID

    Returns:
        LaTeX source
    """
    if graph_service is None:
        raise HTTPException(
            status_code=503,
            detail="API service not initialized. Model backend is unavailable.",
        )
    
    result = _get_task_result(task_id)

    return {"latex_source": result.get("latex_source", "")}


@router.get("/export/json/{task_id}")
async def export_json(task_id: str):
    """
    Export paper as JSON.

    Args:
        task_id: Task ID

    Returns:
        JSON file with all paper data
    """
    if graph_service is None:
        raise HTTPException(
            status_code=503,
            detail="API service not initialized. Model backend is unavailable.",
        )

    result = _get_task_result(task_id)

    with tempfile.NamedTemporaryFile(mode="w", suffix=".json", delete=False) as f:
        json.dump(result, f, indent=2)
        temp_path = f.name

    file_stem = _safe_filename(result.get("topic", "paper"))
    return FileResponse(
        path=temp_path,
        media_type="application/json",
        filename=f"{file_stem}.json",
    )


@router.get("/export/txt/{task_id}")
async def export_txt(task_id: str):
    """
    Export paper as plain text.

    Args:
        task_id: Task ID

    Returns:
        Text file with formatted paper
    """
    if graph_service is None:
        raise HTTPException(
            status_code=503,
            detail="API service not initialized. Model backend is unavailable.",
        )

    result = _get_task_result(task_id)

    lines = []
    lines.append("=" * 80)
    lines.append(result.get("topic", "Research Paper").upper())
    lines.append("=" * 80)
    lines.append("")

    if result.get("keywords"):
        lines.append("Keywords: " + ", ".join(result["keywords"]))
        lines.append("")

    for section in ["abstract", "introduction", "related_work", "methodology", "implementation", "results_discussion", "conclusion"]:
        body = result.get(section) or result.get("results", "") if section == "results_discussion" else result.get(section)
        if body:
            lines.append("\n" + section.upper().replace("_", " "))
            lines.append("-" * 40)
            lines.append(body)

    refs = result.get("references_raw", []) or result.get("references", [])
    if refs:
        lines.append("\nREFERENCES")
        lines.append("-" * 40)
        for i, ref in enumerate(refs, 1):
            label = ref.get("title", str(ref)) if isinstance(ref, dict) else str(ref)
            lines.append(f"[{i}] {label}")

    text_content = "\n".join(lines)

    with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False) as f:
        f.write(text_content)
        temp_path = f.name

    file_stem = _safe_filename(result.get("topic", "paper"))
    return FileResponse(
        path=temp_path,
        media_type="text/plain",
        filename=f"{file_stem}.txt",
    )


@router.get("/export/latex/{task_id}")
async def export_latex(task_id: str):
    """
    Export paper as LaTeX source file.

    Args:
        task_id: Task ID

    Returns:
        LaTeX source file
    """
    if graph_service is None:
        raise HTTPException(
            status_code=503,
            detail="API service not initialized. Model backend is unavailable.",
        )

    result = _get_task_result(task_id)

    latex_source = result.get("latex_source", "")
    if not latex_source:
        raise HTTPException(status_code=404, detail="No LaTeX source available for this task")

    with tempfile.NamedTemporaryFile(mode="w", suffix=".tex", delete=False) as f:
        f.write(latex_source)
        temp_path = f.name

    file_stem = _safe_filename(result.get("topic", "paper"))
    return FileResponse(
        path=temp_path,
        media_type="text/plain",
        filename=f"{file_stem}.tex",
    )


@router.get("/export/docx/{task_id}")
async def export_docx(task_id: str):
    """
    Export paper as DOCX file.

    Args:
        task_id: Task ID

    Returns:
        DOCX file download
    """
    if graph_service is None:
        raise HTTPException(
            status_code=503,
            detail="API service not initialized. Model backend is unavailable.",
        )

    result = _get_task_result(task_id)

    # Use pre-generated DOCX if it exists, otherwise build on-the-fly
    path = result.get("docx_path")
    if not path or not Path(path).exists():
        try:
            path = _build_fallback_docx(result, task_id)
        except Exception as exc:
            raise HTTPException(status_code=500, detail=f"DOCX generation failed: {exc}")

    file_stem = _safe_filename(result.get("topic", "paper"))
    return FileResponse(
        path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"{file_stem}.docx",
    )


@router.get("/review/{task_id}")
async def get_review(task_id: str):
    """Return the ReviewAgent report for a completed task."""
    if graph_service is None:
        raise HTTPException(
            status_code=503,
            detail="API service not initialized. Model backend is unavailable.",
        )

    result = _get_task_result(task_id)

    report = result.get("review_report")
    if not report:
        raise HTTPException(status_code=404, detail="Review report not available.")

    return {
        "task_id":              task_id,
        "review_report":        report,
        "review_score":         result.get("review_score", 0.0),
        "ready_for_submission": result.get("ready_for_submission", False),
    }


@router.get("/tasks", response_model=TaskListResponse)
async def list_tasks() -> TaskListResponse:
    """
    List all paper generation tasks.

    Returns:
        List of tasks with their status
    """
    if graph_service is None:
        raise HTTPException(
            status_code=503,
            detail="API service not initialized. Model backend is unavailable.",
        )
    
    tasks = graph_service.list_tasks()
    return TaskListResponse(tasks=tasks, total=len(tasks))


@router.delete("/tasks/{task_id}")
async def delete_task(task_id: str):
    """
    Delete a task.

    Args:
        task_id: Task ID to delete

    Returns:
        Success message
    """
    if graph_service is None:
        raise HTTPException(
            status_code=503,
            detail="API service not initialized. Model backend is unavailable.",
        )
    
    if graph_service.delete_task(task_id):
        return {"message": f"Task {task_id} deleted"}
    else:
        raise HTTPException(status_code=404, detail="Task not found")


@router.get("/graph-info")
async def get_graph_info():
    """
    Get information about the graph structure.

    Returns:
        Graph structure (nodes, edges, etc.)
    """
    # Create a temporary graph instance to get info
    from backend.services.model_manager import ModelManager
    from backend.graph.paper_graph import PaperGenerationGraph

    graph = PaperGenerationGraph(ModelManager())
    return graph.get_graph_structure()


# ── IEEE Formatting Endpoints ──


@router.get("/ieee-preview/{task_id}")
async def get_ieee_preview(task_id: str):
    """
    Get IEEE-formatted LaTeX preview for a completed paper.

    Args:
        task_id: Task ID

    Returns:
        LaTeX source code in IEEE format
    """
    if graph_service is None:
        raise HTTPException(
            status_code=503,
            detail="API service not initialized. Model backend is unavailable.",
        )

    try:
        result = _get_task_result(task_id)

        from backend.services.ieee_formatter import IEEEPaperFormatter

        formatter = IEEEPaperFormatter()

        # Format the paper using IEEE standards
        formatted = formatter.format_paper(
            title=result.get("topic", result.get("title", "Untitled Paper")),
            authors=_authors_from_result(result),
            abstract=result.get("abstract", ""),
            keywords=result.get("keywords", []),
            sections=_sections_from_result(result),
            equations=result.get("equations", []),
            diagrams=result.get("diagrams", []) or result.get("figures", []),
            references=result.get("references", []),
            tables=result.get("tables", []),
        )

        return {
            "task_id": task_id,
            "latex_source": formatted["latex_source"],
            "sections_count": formatted["sections_count"],
            "equations_count": formatted["equations_count"],
            "diagrams_count": formatted["diagrams_count"],
            "references_count": formatted["references_count"],
        }

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Preview generation failed: {str(e)}")


@router.post("/compile-pdf/{task_id}")
async def compile_pdf(task_id: str):
    """
    Compile an IEEE-formatted paper to PDF.

    Args:
        task_id: Task ID

    Returns:
        PDF file
    """
    if graph_service is None:
        raise HTTPException(
            status_code=503,
            detail="API service not initialized. Model backend is unavailable.",
        )

    try:
        result = _get_task_result(task_id)

        from backend.services.ieee_formatter import IEEEPaperFormatter

        formatter = IEEEPaperFormatter()

        # Get IEEE LaTeX
        formatted = formatter.format_paper(
            title=result.get("topic", result.get("title", "Untitled Paper")),
            authors=_authors_from_result(result),
            abstract=result.get("abstract", ""),
            keywords=result.get("keywords", []),
            sections=_sections_from_result(result),
            equations=result.get("equations", []),
            diagrams=result.get("diagrams", []) or result.get("figures", []),
            references=result.get("references", []),
            tables=result.get("tables", []),
        )

        latex_source = formatted["latex_source"]

        # Try pdflatex; on failure return the .tex for local compilation
        from backend.services.pdf_generator import PDFGenerator

        if PDFGenerator.pdflatex_version():
            pdf_path = PDFGenerator.generate(latex_source, task_id)
        else:
            pdf_path = None

        if pdf_path and Path(pdf_path).exists():
            return FileResponse(
                path=pdf_path,
                filename="paper.pdf",
                media_type="application/pdf",
            )

        # pdflatex missing or compilation failed — return .tex
        tex_out = Path("outputs") / f"{task_id}.tex"
        tex_out.parent.mkdir(parents=True, exist_ok=True)
        tex_out.write_text(latex_source, encoding="utf-8")
        file_stem = _safe_filename(result.get("topic", "paper"))
        return FileResponse(
            path=str(tex_out),
            media_type="text/plain",
            filename=f"{file_stem}_paper.tex",
            headers={"X-Fallback-Format": "latex", "X-Fallback-Reason": "pdflatex-unavailable"},
        )

    except HTTPException:
        raise
    except FileNotFoundError as e:
        raise HTTPException(
            status_code=500,
            detail="pdflatex not found. Install: brew install basictex (macOS)",
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF compilation failed: {str(e)}")


@router.get("/ieee-validate/{task_id}")
async def validate_ieee_format(task_id: str):
    """
    Validate IEEE LaTeX format for a paper.

    Args:
        task_id: Task ID

    Returns:
        Validation results
    """
    if graph_service is None:
        raise HTTPException(
            status_code=503,
            detail="API service not initialized. Model backend is unavailable.",
        )

    try:
        result = _get_task_result(task_id)

        from backend.services.ieee_formatter import IEEEPaperFormatter

        formatter = IEEEPaperFormatter()

        # Format the paper
        formatted = formatter.format_paper(
            title=result.get("title", "Untitled Paper"),
            authors=_authors_from_result(result),
            abstract=result.get("abstract", ""),
            keywords=result.get("keywords", []),
            sections=result.get("sections", {}),
            equations=result.get("equations", []),
            diagrams=result.get("diagrams", []) or result.get("figures", []),
            references=result.get("references", []),
            tables=result.get("tables", []),
        )

        # Validate LaTeX
        validation = formatter.validate_latex(formatted["latex_source"])

        return {
            "task_id": task_id,
            "validation": validation,
            "formatted_info": {
                "sections": formatted["sections_count"],
                "equations": formatted["equations_count"],
                "diagrams": formatted["diagrams_count"],
                "references": formatted["references_count"],
            },
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Validation failed: {str(e)}")


@router.get("/health")
async def health_check():
    """
    Health check endpoint.

    Returns:
        Health status
    """
    from backend.services.pdf_generator import PDFGenerator
    pdflatex_ver = PDFGenerator.pdflatex_version()
    return {
        "status": "ok",
        "service": "langgraph-pipeline",
        "model_provider": model_provider,
        "ollama_base_url": os.getenv("OLLAMA_BASE_URL", "http://localhost:11434"),
        "default_model": default_model,
        "pdflatex": pdflatex_ver or "not installed",
        "pdf_ready": pdflatex_ver is not None,
    }
